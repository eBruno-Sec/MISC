"""
EXCALIBUR  //  Phase 4 — Report Engine
Generates TXT and DOCX reports from all phase results.
"""

import json
from pathlib import Path
from datetime import datetime

R="\033[91m"; Y="\033[93m"; G="\033[92m"; C="\033[96m"; BOLD="\033[1m"; RST="\033[0m"

def ok(m):  print(f"  {G}[+]{RST} {m}")
def info(m):print(f"  {C}[*]{RST} {m}")
def warn(m):print(f"  {Y}[!]{RST} {m}")

SEP1 = "=" * 70
SEP2 = "-" * 70

def build_txt(target, results):
    lines = []
    ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    lines += [
        SEP1,
        "ROUND TABLE  //  BUG BOUNTY INTELLIGENCE REPORT",
        f"Target    : {target}",
        f"Generated : {ts}",
        f"Suite     : Merlin > Percival > Galahad > Lancelot > Excalibur",
        SEP1, "",
    ]

    # ── LANCELOT AI TRIAGE ──
    lance = results.get("lancelot", {})
    if lance.get("triage"):
        lines += [
            "SECTION 1  //  AI TRIAGE (LANCELOT)",
            SEP2,
            f"Model : {lance.get('model','?')}",
            "",
            lance["triage"],
            "",
        ]
    else:
        lines += ["SECTION 1  //  AI TRIAGE: Not available (no API key or error)", ""]

    # ── GAWAIN MANUAL PLAYBOOK ──
    gaw = results.get("gawain", {})
    if gaw.get("playbook"):
        lines += [
            "SECTION 1B  //  MANUAL HUNTING PLAYBOOK (GAWAIN)",
            SEP2,
            f"Model : {gaw.get('model','?')}",
            "Execute one step at a time. The tool sends nothing.",
            "",
            gaw["playbook"],
            "",
        ]

    # ── PERCIVAL ──
    p = results.get("percival", {})
    lines += ["SECTION 2  //  PASSIVE RECON (PERCIVAL)", SEP2, ""]

    w = p.get("whois", {})
    lines += ["[ WHOIS ]"]
    if w:
        lines += [
            f"  Registrar : {w.get('registrar','N/A')}",
            f"  Created   : {w.get('created','N/A')}",
            f"  Expires   : {w.get('expires','N/A')}",
            f"  Updated   : {w.get('updated','N/A')}",
            f"  NS        : {', '.join(w.get('nameservers',[]))}",
            f"  Status    : {', '.join(w.get('status',[]))}",
            f"  Privacy   : {'Redacted' if w.get('privacy_redacted') else 'Not redacted'}",
        ]
    else:
        lines.append("  Unavailable")
    lines.append("")

    lines += ["[ DNS RECORDS ]"]
    for label, key in [("A","a_records"),("AAAA","aaaa_records"),("MX","mx_records"),("NS","ns_records"),("TXT","txt_records"),("CAA","caa_records")]:
        recs = p.get(key, [])
        lines.append(f"  {label:6}: {', '.join(recs) if recs else 'None'}")
    lines.append("")

    em = p.get("email", {})
    lines += [
        "[ EMAIL SECURITY ]",
        f"  SPF  : {em.get('spf','MISSING')}",
        f"  DMARC: {em.get('dmarc','MISSING')}",
        f"  BIMI : {em.get('bimi','Not configured')}",
        f"  DKIM : {[d['selector'] for d in em.get('dkim',[])] or 'None found'}",
        "",
    ]

    http = p.get("http", {})
    lines += ["[ HTTP SECURITY HEADERS ]"]
    if http.get("ok"):
        h = http.get("headers", {})
        lines.append(f"  Status: {http.get('status')}  HTTPS: {http.get('is_https')}")
        for name in ["strict-transport-security","content-security-policy","x-frame-options",
                     "x-content-type-options","referrer-policy","permissions-policy"]:
            present = name in h
            lines.append(f"  {'[OK]    ' if present else '[MISSING]'} {name}")
        if h.get("server"):         lines.append(f"  Server      : {h['server']}")
        if h.get("x-powered-by"):   lines.append(f"  X-Powered-By: {h['x-powered-by']}")
    else:
        lines.append("  Header inspection unavailable")
    lines.append("")

    ssl_d = p.get("ssl", {})
    lines += ["[ SSL CERTIFICATE ]"]
    if "error" not in ssl_d:
        lines += [
            f"  Issued To  : {ssl_d.get('issued_to','N/A')}",
            f"  Issued By  : {ssl_d.get('issued_by','N/A')}",
            f"  Valid Until: {ssl_d.get('valid_until','N/A')}",
            f"  Days Left  : {ssl_d.get('days_left','N/A')}",
            f"  SANs       : {', '.join(ssl_d.get('san',[])[:5])}",
        ]
    else:
        lines.append(f"  Error: {ssl_d.get('error')}")
    lines.append("")

    vendors = p.get("vendors", [])
    lines += ["[ TECH STACK ]"]
    for v in vendors:
        lines.append(f"  [{v['rv']:8}] {v['name']:30} {v['cat']}  (via {v.get('source','?')})")
    if not vendors:
        lines.append("  No vendors detected")
    lines.append("")

    sub_cats = p.get("sub_cats", {})
    subs     = p.get("subdomains", [])
    wildcards= p.get("wildcards", [])
    lines += [f"[ SUBDOMAINS ({len(subs)} found, {len(wildcards)} wildcards) ]"]
    if wildcards:
        lines.append(f"  WILDCARDS: {', '.join(wildcards[:5])}")
    priority_cats = ["CI/CD & DevOps","Security Infrastructure","Admin & Management","Payment & Financial","Exposed Dev/Test"]
    for cat in priority_cats:
        if cat in sub_cats:
            cat_subs = [s["name"] for s in sub_cats[cat]]
            lines.append(f"  [{cat}] ({len(cat_subs)} subdomains)")
            for s in cat_subs[:10]:
                lines.append(f"    -> {s}")
    lines.append("")

    # ── GALAHAD ──
    g = results.get("galahad", {})
    if g:
        lines += ["SECTION 3  //  ACTIVE ENUMERATION (GALAHAD)", SEP2, ""]

        live = g.get("live_hosts", [])
        lines += [f"[ LIVE HOSTS ({len(live)}) ]"]
        for h in live[:20]:
            lines.append(f"  {h.get('url','')}  [{h.get('status-code','?')}]  {','.join(h.get('tech',[])[:3])}")
        if len(live) > 20:
            lines.append(f"  ... and {len(live)-20} more")
        lines.append("")

        nmap = g.get("nmap", {})
        lines += ["[ PORT SCAN ]"]
        for port_line in nmap.get("open_ports", [])[:20]:
            lines.append(f"  {port_line.strip()}")
        if not nmap.get("open_ports"):
            lines.append("  No open ports found or nmap unavailable")
        lines.append("")

        nuclei = g.get("nuclei", [])
        lines += [f"[ NUCLEI FINDINGS ({len(nuclei)}) ]"]
        for n in nuclei[:20]:
            inf = n.get("info", {})
            sev = inf.get("severity","?")
            nm  = inf.get("name","") or n.get("raw","")[:60]
            host= n.get("host","")
            lines.append(f"  [{sev.upper():8}] {nm}  {host}")
        if not nuclei:
            lines.append("  No findings or nuclei unavailable")
        lines.append("")

        misc = g.get("misc", [])
        lines += [f"[ CORS / VCS FINDINGS ({len(misc)}) ]"]
        for m in misc:
            lines.append(f"  [{m.get('severity','?'):8}] {m.get('type','?')}  {m.get('url','')}  {m.get('detail','')}")
        if not misc:
            lines.append("  None found")
        lines.append("")

        takeovers = g.get("takeover_candidates", [])
        crit_to = [t for t in takeovers if t.get("severity") == "CRITICAL"]
        lines += [f"[ SUBDOMAIN TAKEOVER CANDIDATES ({len(crit_to)} critical) ]"]
        for t in crit_to:
            lines.append(f"  [CRITICAL] {t.get('subdomain','')}  {t.get('reason','')}")
        if not crit_to:
            lines.append("  No critical takeover candidates found")
        lines.append("")

        dir_bust = g.get("dir_bust", {})
        lines += ["[ DIRECTORY BUSTING ]"]
        for url, paths in dir_bust.items():
            if paths:
                lines.append(f"  {url}: {len(paths)} paths")
                for pth in paths[:5]:
                    lines.append(f"    {pth.get('url','')} [{pth.get('status','')}]")
        if not dir_bust:
            lines.append("  No paths found or ffuf/gobuster unavailable")
        lines.append("")

    lines += [
        SEP1,
        "PASSIVE SCOPE: Percival uses only public data (DNS, crt.sh, RDAP, HTTP headers)",
        "ACTIVE SCOPE:  Galahad sends packets to target. Only use within authorized scope.",
        "Round Table by Erwin Bruno  //  github.com/eBruno-Sec",
        SEP1,
    ]
    return "\n".join(lines)

def build_docx(target, results, out_path):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        warn("python-docx not installed. Skipping DOCX export.")
        return

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Consolas"
    style.font.size = Pt(9)

    def h1(text):
        p = doc.add_heading(text, level=1)
        p.runs[0].font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
        return p

    def h2(text):
        return doc.add_heading(text, level=2)

    def mono(text, color=None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        if color:
            run.font.color.rgb = color
        return p

    ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    h1("ROUND TABLE  //  BUG BOUNTY INTELLIGENCE REPORT")
    mono(f"Target    : {target}")
    mono(f"Generated : {ts}")
    mono("Suite     : Merlin > Percival > Galahad > Lancelot > Excalibur")
    doc.add_paragraph()

    # AI Triage
    lance = results.get("lancelot", {})
    h2("AI TRIAGE (LANCELOT)")
    if lance.get("triage"):
        mono(f"Model: {lance.get('model','?')}")
        doc.add_paragraph()
        for line in lance["triage"].split("\n"):
            mono(line)
    else:
        mono("AI triage not available (no API key or error)")

    # Gawain manual playbook
    gaw = results.get("gawain", {})
    if gaw.get("playbook"):
        h2("MANUAL HUNTING PLAYBOOK (GAWAIN)")
        mono(f"Model: {gaw.get('model','?')}")
        mono("Execute one step at a time. The tool sends nothing.")
        doc.add_paragraph()
        for line in gaw["playbook"].split("\n"):
            mono(line)

    # Percival
    p = results.get("percival", {})
    h2("PASSIVE RECON (PERCIVAL)")

    w = p.get("whois", {})
    h2("WHOIS")
    if w:
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Field"
        hdr_cells[1].text = "Value"
        for field, val in [("Registrar",w.get("registrar","N/A")),("Created",w.get("created","N/A")),
                           ("Expires",w.get("expires","N/A")),("Privacy","Yes" if w.get("privacy_redacted") else "No")]:
            row = table.add_row().cells
            row[0].text = field
            row[1].text = str(val)
        doc.add_paragraph()

    h2("TECH STACK")
    vendors = p.get("vendors", [])
    if vendors:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        for i, hd in enumerate(["Vendor","Category","Recon Value","Source"]):
            table.rows[0].cells[i].text = hd
        for v in vendors:
            row = table.add_row().cells
            row[0].text = v.get("name","")
            row[1].text = v.get("cat","")
            row[2].text = v.get("rv","")
            row[3].text = v.get("source","")
        doc.add_paragraph()

    # Galahad
    g = results.get("galahad", {})
    if g:
        h2("ACTIVE ENUMERATION (GALAHAD)")
        nuclei = g.get("nuclei", [])
        if nuclei:
            h2("Nuclei Findings")
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            for i, hd in enumerate(["Severity","Finding","Host"]):
                table.rows[0].cells[i].text = hd
            for n in nuclei[:30]:
                inf = n.get("info", {})
                row = table.add_row().cells
                row[0].text = inf.get("severity","?").upper()
                row[1].text = inf.get("name","") or n.get("raw","")[:60]
                row[2].text = n.get("host","")
            doc.add_paragraph()

        takeovers = [t for t in g.get("takeover_candidates",[]) if t.get("severity")=="CRITICAL"]
        if takeovers:
            h2("Subdomain Takeover Candidates (CRITICAL)")
            for t in takeovers:
                mono(f"[CRITICAL] {t.get('subdomain','')}  {t.get('reason','')}")

    doc.save(out_path)
    ok(f"DOCX report saved: {out_path}")

# ─── MAIN ──────────────────────────────────────────────────────────────────────
def run_excalibur(target, results, run_dir, cfg):
    run_dir = Path(run_dir)

    # TXT report
    info("Generating TXT report...")
    txt_content = build_txt(target, results)
    txt_path = run_dir / f"report_{target}.txt"
    txt_path.write_text(txt_content, encoding="utf-8")
    ok(f"TXT report: {txt_path}")

    # DOCX report
    info("Generating DOCX report...")
    docx_path = run_dir / f"report_{target}.docx"
    build_docx(target, results, str(docx_path))

    # JSON dump (all raw data)
    info("Saving full JSON dump...")
    json_path = run_dir / f"report_{target}_raw.json"
    json_path.write_text(json.dumps(results, indent=2, default=str))
    ok(f"JSON dump: {json_path}")

    # Summary to stdout
    p = results.get("percival", {})
    g = results.get("galahad", {})
    lance = results.get("lancelot", {})

    print(f"\n{G}{BOLD}EXCALIBUR REPORT SUMMARY{RST}")
    print(f"  Target           : {target}")
    print(f"  Subdomains (P1)  : {len(p.get('subdomains',[]))}")
    print(f"  Subdomains (P2)  : {len(g.get('all_subdomains',[]))}" if g else "  Phase 2         : skipped")
    if g:
        print(f"  Live Hosts       : {len(g.get('live_hosts',[]))}")
        print(f"  Nuclei Findings  : {len(g.get('nuclei',[]))}")
        print(f"  Takeover Cands   : {len([t for t in g.get('takeover_candidates',[]) if t.get('severity')=='CRITICAL'])}")
        print(f"  CORS/VCS Issues  : {len(g.get('misc',[]))}")
    print(f"  AI Triage        : {'complete' if lance.get('triage') else 'skipped'}")
    gaw = results.get("gawain", {})
    print(f"  Manual Playbook  : {'complete' if gaw.get('playbook') else 'skipped'}")
    print(f"  Reports          : {txt_path.name} | {docx_path.name}")
    print(f"  Output Dir       : {run_dir}")
